"""Export application packages to DOCX, PDF, and TXT."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from application_workspace.formatting import format_interview_prep, format_quality
from application_workspace.models import ApplicationPackage


def _section(title: str, body: str) -> str:
    return f"\n{'=' * 60}\n{title}\n{'=' * 60}\n\n{body.strip()}\n"


def package_as_plaintext(package: ApplicationPackage) -> str:
    """Single plaintext bundle for export."""
    snap = package.recommendation_snapshot or {}
    fit = snap.get("overall_match")
    conf = snap.get("confidence")
    score_lines: list[str] = []
    if fit is not None and float(fit) > 0:
        score_lines.append(f"Fit: {int(round(float(fit) * 100))}%")
    if conf is not None and float(conf) > 0:
        score_lines.append(f"Confidence: {int(round(float(conf) * 100))}%")
    score_block = " · ".join(score_lines)

    parts = [
        f"APPLICATION PACKAGE — {package.job_title} @ {package.company}",
        f"Package ID: {package.package_id}",
        f"Job ID: {package.source_job_id}",
        f"Status: {package.approval_status.value}",
        f"Generated: {package.generated_at}",
    ]
    if score_block:
        parts.append(f"Match scores: {score_block}")
    parts.extend(
        [
            "",
            "EXPLAINABILITY",
            *[f"  • {e}" for e in package.explanation],
        ],
    )
    if package.tailored_resume_text:
        parts.append(_section("TAILORED RESUME", package.tailored_resume_text))
    if package.cover_letter:
        parts.append(_section("COVER LETTER", package.cover_letter.body))
    if package.recruiter_message:
        parts.append(_section("LINKEDIN INTRO", package.recruiter_message.linkedin_intro))
        parts.append(_section("HIRING MANAGER NOTE", package.recruiter_message.hiring_manager_note))
        parts.append(_section("REFERRAL REQUEST", package.recruiter_message.referral_request))
    if package.interview_prep:
        parts.append(_section("INTERVIEW PREP", format_interview_prep(package.interview_prep)))
    if package.quality_scores:
        parts.append(_section("QUALITY SCORES", format_quality(package.quality_scores)))
    return "\n".join(parts)


def export_txt(package: ApplicationPackage, out_dir: Path) -> Path:
    from demo_mode import persistence_writes_enabled

    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{package.package_id}_application_package.txt"
    if not persistence_writes_enabled():
        return path
    path.write_text(package_as_plaintext(package), encoding="utf-8")
    return path


def export_docx(package: ApplicationPackage, out_dir: Path) -> Path:
    from docx import Document

    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{package.package_id}_application_package.docx"
    doc = Document()
    doc.add_heading(f"Application Package — {package.job_title}", level=1)
    doc.add_paragraph(f"{package.company} · {package.approval_status.value}")
    doc.add_heading("Explainability", level=2)
    for line in package.explanation:
        doc.add_paragraph(line, style="List Bullet")
    if package.cover_letter:
        doc.add_heading("Cover Letter", level=2)
        for para in package.cover_letter.body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    if package.recruiter_message:
        doc.add_heading("Recruiter Messages", level=2)
        doc.add_paragraph(package.recruiter_message.linkedin_intro)
    if package.tailored_resume_text:
        doc.add_heading("Tailored Resume", level=2)
        for line in package.tailored_resume_text.splitlines():
            if line.strip():
                doc.add_paragraph(line.strip())
    doc.save(str(path))
    return path


def export_pdf(package: ApplicationPackage, out_dir: Path) -> Path:
    """PDF via reportlab simple flow; falls back to TXT if reportlab missing."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        return export_txt(package, out_dir)

    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{package.package_id}_application_package.pdf"
    text = package_as_plaintext(package)
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 50
    line_height = 12
    for raw_line in text.splitlines():
        line = raw_line[:100]
        if y < 50:
            c.showPage()
            y = height - 50
        c.drawString(50, y, line)
        y -= line_height
    c.save()
    return path


def export_package(
    package: ApplicationPackage,
    export_dir: Path,
    formats: list[str] | None = None,
) -> dict[str, Any]:
    """Export package artifacts; returns paths by format."""
    formats = formats or ["txt", "docx", "pdf"]
    out = export_dir / package.package_id
    paths: dict[str, str] = {}
    if "txt" in formats:
        paths["txt"] = str(export_txt(package, out))
    if "docx" in formats:
        paths["docx"] = str(export_docx(package, out))
    if "pdf" in formats:
        paths["pdf"] = str(export_pdf(package, out))
    return paths
